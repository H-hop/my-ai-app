import os
from docx import Document

# --- 这里是关键：请确认你的 Word 文件夹路径 ---
# 建议在桌面新建一个文件夹叫 words，把要处理的文件都放进去
FOLDER_PATH = r'C:\Users\Administrator\Desktop\words' 
OUTPUT_NAME = '整合后的总文件.docx'
# ------------------------------------------

def combine_words():
    # 1. 检查文件夹是否存在
    if not os.path.exists(FOLDER_PATH):
        print(f"❌ 找不到文件夹: {FOLDER_PATH}")
        print("💡 请手动在桌面创建一个名为 words 的文件夹，并把Word放进去。")
        return

    # 2. 获取所有 docx 文件并按学号（文件名）排序
    files = [f for f in os.listdir(FOLDER_PATH) if f.endswith('.docx') and not f.startswith('~$')]
    files.sort() # 这里会自动按文件名（学号）从小到大排列
    
    if not files:
        print("📂 文件夹里是空的，没找到 .docx 文件。")
        return

    print(f"✅ 找到 {len(files)} 个文件，准备按此顺序合并：{files}")

    main_doc = Document()

    for index, filename in enumerate(files):
        file_path = os.path.join(FOLDER_PATH, filename)
        sub_doc = Document(file_path)
        
        print(f"正在处理: {filename}...")

        if index == 0:
            # 学号第一的：完整保留（直接把所有段落搬过去）
            for para in sub_doc.paragraphs:
                new_para = main_doc.add_paragraph(para.text)
                new_para.style = para.style
        else:
            # 学号之后的：删除第一页
            # 逻辑：寻找第一个分页符，只保留分页符之后的内容
            found_page_break = False
            for para in sub_doc.paragraphs:
                # 检查段落里是否有分页符（Manual Page Break）
                if 'w:br' in para._element.xml and 'type="page"' in para._element.xml:
                    found_page_break = True
                    continue 
                
                if found_page_break:
                    # 只有在发现第一个分页符之后，才开始添加内容
                    new_para = main_doc.add_paragraph(para.text)
                    new_para.style = para.style

        # 每个人的文档合并完，加一个分页符隔开
        main_doc.add_page_break()

    # 4. 保存文件
    main_doc.save(OUTPUT_NAME)
    print(f"\n✨ 处理完成！文件已生成在当前目录下：{OUTPUT_NAME}")

if __name__ == "__main__":
    combine_words()
